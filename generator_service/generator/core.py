import time
import logging
import aiohttp
from docx import Document

logger = logging.getLogger(__name__)


def generate_fake_document(user_data: dict, content_type: str) -> str:
    user_name = user_data.get("first_name", "N/A")
    user_id = user_data.get("id", "N/A")
    logger.info(f"Начал генерацию {content_type} для {user_name}...")

    if content_type == "docx":
        doc = Document()
        doc.add_heading(f'Карточка:{user_data.get("last_name")}', 0)
        doc.add_paragraph(f'ИИН:{user_data.get("iin")}')
        doc.add_paragraph(f'Номер телефона:{user_data.get("phone_number")}')
    else:
        time.sleep(3)

    fake_filename = f"user_{user_id}_profile.{content_type}"
    fake_url = f"/generated_docs/{fake_filename}"

    logger.info(f"Завершил генерацию для {user_name}. Ссылка: {fake_url}")
    return fake_url


async def send_callback(callback_url: str, payload: dict):
    logger.info(f"Отправляю POST на {callback_url}...")
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(callback_url, json=payload) as response:
                if response.status == 200:
                    logger.info(f"Callback на {callback_url} успешно доставлен")
                else:
                    logger.warning(
                        f"Ошибка доставки callback, статус: {response.status}"
                    )
    except Exception as e:
        logger.error(f"КРИТИЧЕСКАЯ ОШИБКА отправки callback {e}")
